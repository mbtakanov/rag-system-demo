import os
import re
import sys
import random
import logging
import unicodedata
from time import sleep
from datetime import datetime
import xml.etree.ElementTree as ET

import requests
from tqdm import tqdm
from docx import Document
from openai import OpenAI
from PyPDF2 import PdfReader
from dotenv import load_dotenv

from src.config import DOCX_DIR, PDF_DIR

load_dotenv()
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(name)s: %(message)s",
    handlers=[logging.StreamHandler(sys.stdout)],
)
logger = logging.getLogger(__name__)
client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

os.makedirs(DOCX_DIR, exist_ok=True)
os.makedirs(PDF_DIR, exist_ok=True)

TOPICS = [
    "AI in Education",
    "Machine Learning in Retail",
    "Natural Language Processing in Law",
    "Ethical Data Collection",
    "Deep Learning in Climate Science",
    "Human-AI Collaboration",
    "Edge Computing for IoT",
    "Predictive Maintenance",
    "Robotics in Manufacturing",
    "AI in Marketing Analytics",
    "Facial Recognition Regulations",
    "Smart Cities and AI",
    "Healthcare Predictive Analytics",
    "Algorithmic Bias in Hiring",
    "Speech Recognition Trends",
    "AI in Logistics",
    "Computer Vision in Sports",
    "Neural Networks in Music",
    "Chatbots in Customer Support",
    "AI in Agriculture",
    "AI for Legal Research",
    "AI in Mental Health",
    "Computer Vision in Retail",
    "AI for Disaster Management",
    "AI in Sports Analytics",
    "AI-driven HR Systems",
    "AI for Wildlife Conservation",
    "AI in Urban Planning",
    "Voice Recognition in Healthcare",
    "AI and Cultural Heritage",
    "AI in Manufacturing Efficiency",
    "AI and Creativity",
    "AI in Transportation",
    "AI in Marketing Forecasting",
    "AI for Fraud Detection",
    "AI in Cybersecurity Training",
    "AI in Language Translation",
    "AI for Traffic Management",
    "AI in Customer Experience",
    "AI and Human Behavior",
    "AI in Predictive Policing",
    "AI in Drug Development",
    "AI in eCommerce Personalization",
    "AI in Journalism",
    "AI for Environmental Monitoring",
    "AI in Retail Supply Chain",
    "AI for Financial Forecasting",
    "AI in Energy Efficiency",
    "AI and Data Privacy",
    "AI in Space Exploration",
    "AI for Smart Homes",
    "AI for Small Businesses",
    "AI in Architecture",
    "AI in Fashion Design",
    "AI in Insurance",
    "AI in Gaming",
    "AI in Education Policy",
    "AI in Music Creation",
    "AI and the Future of Work",
    "AI in Neuroscience",
    "AI for Social Good",
    "AI for Accessibility",
    "AI in Automotive Safety",
    "AI and Ethics in Law",
    "AI in Art Restoration",
    "AI in Personalized Medicine",
    "AI in Climate Adaptation",
    "AI and Governance",
    "AI in Public Policy",
    "AI for Scientific Discovery",
    "AI and Creativity in Media",
    "AI in Construction Management",
    "AI in Retail Analytics",
    "AI for Mental Wellness",
    "AI in Government Services",
]


def generate_report(topic: str) -> str:
    # The query can be improved to be more specific, detailed, add tables, images, etc.
    response = client.responses.create(
        model="gpt-4.1-nano",
        input=[
            {
                "role": "system",
                "content": "You are a professional research report writer.",
            },
            {
                "role": "user",
                "content": (
                    f"Write a detailed research report of at least 3000 words on '{topic}'. "
                    "Write only plain text - do not include any figures, images, captions, diagrams, or references "
                    "(e.g., '(Image: ...)', '(Figure 1)', or similar). "
                    "Avoid metadata, citations, and formatting markup. "
                    "Use section headings and paragraphs only. "
                    "Do NOT include a summary, conclusion, or recommendations."
                ),
            },
        ],
        max_output_tokens=4000,
    )
    return response.output[0].content[0].text


def generate_docx_set(n_docs=50):
    for i in tqdm(range(n_docs)):
        topic = random.choice(TOPICS)
        try:
            text = generate_report(topic)
            doc = Document()

            doc.core_properties.subject = topic
            doc.core_properties.created = datetime.now()
            doc.add_heading(f"Report on {topic}", level=1)
            doc.add_paragraph(text)

            filename = f"doc_{i+1}_{sanitized_title(topic)}.docx"
            doc.save(os.path.join(DOCX_DIR, filename))
            sleep(1)
        except Exception as e:
            logger.exception(f"Error (DOCX {topic}): {e}")
    logger.info("DOCX generation complete.\n")


def download_arxiv_pdfs(n_pdfs=50, min_pages=5, max_pages=20):
    api_url = "http://export.arxiv.org/api/query"

    downloaded = 0
    start = 0
    batch_size = 100

    while downloaded < n_pdfs:
        params = {
            "search_query": (
                "cat:cs.ML OR cat:cs.LG OR "  # Machine Learning
                "cat:cs.AI OR "  # Artificial Intelligence
                "cat:cs.NE OR "  # Neural & Evolutionary Computing
                "cat:cs.CV OR "  # Computer Vision
                "cat:cs.CL OR "  # Computation & Language (NLP)
                "cat:cs.IR OR "  # Information Retrieval
                "cat:cs.RO OR "  # Robotics
                "cat:cs.HC OR "  # Human-Computer Interaction
                "cat:cs.CR OR "  # Cryptography & Security
                "cat:cs.DB OR "  # Databases
                "cat:stat.ML OR "  # Statistics - Machine Learning
                "cat:q-bio.QM OR "  # Quantitative Methods
                "cat:q-bio.NC OR "  # Neurons & Cognition
                "cat:econ.EM OR "  # Econometrics
                "cat:physics.data-an"  # Data Analysis
            ),
            "start": start,
            "max_results": batch_size,
            "sortBy": "submittedDate",
            "sortOrder": "descending",
        }

        response = requests.get(api_url, params=params)

        if response.status_code != 200:
            logger.error(f"API error: {response.status_code}")
            sleep(5)
            continue

        try:
            root = ET.fromstring(response.content)
        except ET.ParseError as e:
            logger.error(f"XML parse error: {e}")
            logger.error(f"Response: {response.content[:200]}")
            sleep(5)
            continue
        entries = root.findall("{http://www.w3.org/2005/Atom}entry")

        if not entries:
            logger.warning("No more results available, retrying next batch...")
            start += batch_size
            continue

        for entry in tqdm(entries):
            if downloaded >= n_pdfs:
                break

            pdf_link = (
                entry.find("{http://www.w3.org/2005/Atom}id").text.replace(
                    "/abs/", "/pdf/"
                )
                + ".pdf"
            )
            title = entry.find("{http://www.w3.org/2005/Atom}title").text.strip()

            try:
                r = requests.get(pdf_link, timeout=30)
                if r.status_code == 200:
                    temp_path = os.path.join(PDF_DIR, f"temp.pdf")
                    with open(temp_path, "wb") as f:
                        f.write(r.content)

                    reader = PdfReader(temp_path)
                    num_pages = len(reader.pages)

                    if min_pages <= num_pages <= max_pages:
                        safe_title = sanitized_title(title)
                        filename = f"doc_{downloaded+1}_{safe_title}.pdf"
                        final_path = os.path.join(PDF_DIR, filename)
                        os.rename(temp_path, final_path)
                        logger.info(f"Downloaded: {final_path} ({downloaded+1})")
                        downloaded += 1
                    else:
                        os.remove(temp_path)
                        logger.info(f"Skipped: {num_pages} pages")
            except Exception as e:
                logger.exception(f"Error: {e}")

            sleep(3)

        start += batch_size
        sleep(3)

    logger.info(f"Downloaded {downloaded} PDFs.")


def sanitized_title(title: str, max_length: int = 50) -> str:
    """Convert title to safe filename, handling special chars and unicode."""
    title = unicodedata.normalize("NFKD", title)
    title = title.encode("ascii", "ignore").decode("ascii")
    title = re.sub(r"[^\w\s-]", "", title)
    title = re.sub(r"\s+", "_", title.strip())
    title = re.sub(r"_+", "_", title)
    title = re.sub(r"-+", "_", title)
    return title.lower()[:max_length].strip("_")


if __name__ == "__main__":
    generate_docx_set()
    download_arxiv_pdfs()
